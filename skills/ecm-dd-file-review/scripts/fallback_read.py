#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
fallback_read.py — 外部 skill（pdf / docx / xlsx）缺失时的纯 Python 兜底抽取

**优先使用 Anthropic 官方的 `pdf` / `docx` / `xlsx` 外部 skill**，兼容性更好。
本脚本仅在外部 skill 不可用时作为最低限度的兜底：把文件转成纯文本，让 Skill 至少能做文本级别的审阅。

兜底策略：
    - PDF → pypdf 抽正文（扫描件不走 OCR，只能读到"（空）"）
    - DOCX → python-docx 抽段落和表格
    - XLSX → openpyxl 抽每个工作表的 cell 矩阵

依赖：
    pip install -r scripts/requirements.txt

用法：
    python skills/ecm-dd-file-review/scripts/fallback_read.py <file>
    python skills/ecm-dd-file-review/scripts/fallback_read.py <file> --max-chars 5000
"""
from __future__ import annotations

import argparse
import sys
from pathlib import Path


def read_pdf(path: Path) -> str:
    try:
        from pypdf import PdfReader
    except ImportError:
        return "[fallback_read:pdf] 未安装 pypdf；pip install pypdf"
    try:
        reader = PdfReader(str(path))
    except Exception as exc:  # noqa: BLE001
        return f"[fallback_read:pdf] 打开失败：{exc}"
    out: list[str] = []
    for i, page in enumerate(reader.pages, 1):
        try:
            text = page.extract_text() or ""
        except Exception as exc:  # noqa: BLE001
            text = f"[页 {i} 抽取失败：{exc}]"
        out.append(f"\n--- Page {i} ---\n{text}")
    return "\n".join(out).strip() or "（空；可能为扫描件，需走 OCR）"


def read_docx(path: Path) -> str:
    try:
        import docx  # python-docx
    except ImportError:
        return "[fallback_read:docx] 未安装 python-docx；pip install python-docx"
    try:
        doc = docx.Document(str(path))
    except Exception as exc:  # noqa: BLE001
        return f"[fallback_read:docx] 打开失败：{exc}"
    parts: list[str] = []
    # 段落
    for para in doc.paragraphs:
        if para.text.strip():
            style = para.style.name if para.style else ""
            prefix = f"[{style}] " if style and style != "Normal" else ""
            parts.append(f"{prefix}{para.text}")
    # 表格
    for ti, table in enumerate(doc.tables, 1):
        parts.append(f"\n--- Table {ti} ---")
        for row in table.rows:
            cells = [c.text.replace("\n", " / ").strip() for c in row.cells]
            parts.append(" | ".join(cells))
    return "\n".join(parts).strip() or "（空）"


def read_xlsx(path: Path) -> str:
    try:
        from openpyxl import load_workbook
    except ImportError:
        return "[fallback_read:xlsx] 未安装 openpyxl；pip install openpyxl"
    try:
        wb = load_workbook(str(path), data_only=True, read_only=True)
    except Exception as exc:  # noqa: BLE001
        return f"[fallback_read:xlsx] 打开失败：{exc}"
    parts: list[str] = []
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        parts.append(f"\n--- Sheet: {sheet_name} ---")
        for row in ws.iter_rows(values_only=True):
            non_empty = [str(v) if v is not None else "" for v in row]
            if any(non_empty):
                parts.append(" | ".join(non_empty))
    return "\n".join(parts).strip() or "（空）"


def read_plain(path: Path) -> str:
    try:
        return path.read_text(encoding="utf-8", errors="replace")
    except Exception as exc:  # noqa: BLE001
        return f"[fallback_read:plain] 读取失败：{exc}"


READERS = {
    ".pdf": read_pdf,
    ".docx": read_docx,
    ".xlsx": read_xlsx,
    ".txt": read_plain,
    ".md": read_plain,
    ".csv": read_plain,
}


def main(argv: list[str] | None = None) -> int:
    p = argparse.ArgumentParser(
        description="外部 skill 缺失时的最低限度文本抽取兜底工具。",
    )
    p.add_argument("file", help="文件路径")
    p.add_argument(
        "--max-chars",
        type=int,
        default=0,
        help="输出最大字符数，0 表示不截断（默认不截断）",
    )
    args = p.parse_args(argv)

    path = Path(args.file).expanduser().resolve()
    if not path.exists():
        sys.stderr.write(f"错误：文件不存在 {path}\n")
        return 2
    ext = path.suffix.lower()
    if ext not in READERS:
        sys.stderr.write(
            f"错误：不支持的扩展名 {ext}；"
            f"支持：{', '.join(sorted(READERS))}\n"
        )
        return 2

    text = READERS[ext](path)
    if args.max_chars > 0 and len(text) > args.max_chars:
        text = text[: args.max_chars] + f"\n\n... [已截断，总长 {len(text)} 字符]"
    print(text)
    return 0


if __name__ == "__main__":
    sys.exit(main())
